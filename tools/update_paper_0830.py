from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE = Path("paper/S4D_TAM_draft_0830.docx")
OUTPUT = Path("paper/S4D_TAM_draft_0830_updated.docx")


def paragraphs(doc: Document) -> list[Paragraph]:
    return list(doc.paragraphs)


def find_all(doc: Document, text: str) -> list[Paragraph]:
    return [p for p in doc.paragraphs if p.text.strip() == text]


def find_first(doc: Document, text: str) -> Paragraph:
    values = find_all(doc, text)
    if not values:
        raise ValueError(f"Paragraph not found: {text}")
    return values[0]


def find_last(doc: Document, text: str) -> Paragraph:
    values = find_all(doc, text)
    if not values:
        raise ValueError(f"Paragraph not found: {text}")
    return values[-1]


def insert_before(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_after(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def remove_between(start: Paragraph, end: Paragraph) -> None:
    current = start._p.getnext()
    while current is not None and current is not end._p:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt


def insert_blocks_before(anchor: Paragraph, blocks: list[tuple[str, str | None]]) -> None:
    for text, style in blocks:
        insert_before(anchor, text, style)


def replace_exact(doc: Document, old: str, new: str, occurrence: str = "last") -> None:
    p = find_first(doc, old) if occurrence == "first" else find_last(doc, old)
    p.text = new


def replace_contains(doc: Document, needle: str, new: str) -> None:
    for p in doc.paragraphs:
        if needle in p.text:
            p.text = new
            return
    raise ValueError(f"Text fragment not found: {needle}")


def add_section_before(doc: Document, before_heading: str, heading: str, body: list[str]) -> None:
    anchor = find_last(doc, before_heading)
    insert_before(anchor, heading, "Heading 2")
    for text in body:
        insert_before(anchor, text, "Normal")


def replace_section(
    doc: Document,
    start_heading: str,
    end_heading: str,
    blocks: list[tuple[str, str | None]],
) -> None:
    start = find_last(doc, start_heading)
    end = find_last(doc, end_heading)
    remove_between(start, end)
    insert_blocks_before(end, blocks)


def main() -> None:
    doc = Document(SOURCE)

    # Static table of contents updates.
    toc_42 = find_first(doc, "4.2. System Architecture")
    insert_after(toc_42, "4.3. Current Reference Implementation and Target Research Modules")
    replace_exact(doc, "7.2. System Variants", "7.2. Two-Level Comparison Design", "first")
    toc_72 = find_first(doc, "7.2. Two-Level Comparison Design")
    p = insert_after(toc_72, "7.2.1. External System Comparison")
    insert_after(p, "7.2.2. Internal Mechanism Study")
    toc_73 = find_first(doc, "7.3. Metrics")
    insert_after(toc_73, "7.4. Confirmatory Protocol and Statistical Analysis")

    # Abstract and introduction now reflect the executable research reference.
    for p in doc.paragraphs:
        if p.text.startswith("We propose S4D-TAM") and "The article presents" in p.text:
            p.text = p.text.replace(
                "The article presents the system concept, literature review, mathematical description of the architecture, and a proposal for experimental validation methodology.",
                "The article presents the system concept, literature review, mathematical description, an executable auditable reference implementation, and a two-level experimental methodology separating external system comparison from internal mechanism ablation. Confirmatory experimental results are not yet reported.",
            )
            break

    replace_contains(
        doc,
        "The article is conceptual and theoretical in character",
        "The article combines a conceptual system architecture with an executable research reference and a reproducible benchmark framework. The current repository implements token state and lifecycle management, proposal and association, multimodal encoder interfaces, attention-related processing, uncertainty calibration utilities, reference-map and topological support, causal occupancy and motion forecasting, deterministic risk-aware planning, telemetry, and the integrated S4DTAMReference execution path. The final learned hierarchical model, full public-dataset reproduction, and confirmatory results remain under development. Section 7 therefore specifies a prospective two-level validation protocol rather than reporting final performance results.",
    )

    # Implementation-aware architecture section.
    add_section_before(
        doc,
        "5. Mathematical Description",
        "4.3. Current Reference Implementation and Target Research Modules",
        [
            "The S4D-TAM concept is accompanied by an executable Python research reference. The implementation is intentionally modular so that individual mechanisms can be tested, replaced, and ablated without changing the evaluator boundary.",
            "Implemented reference components include persistent token state and lifecycle management (token.py, memory.py), token proposal and data association (proposal.py, association.py), modality-oriented encoder interfaces and masked fusion (encoders/), attention-related processing (attention.py), calibration and uncertainty utilities (calibration.py), reference-map and topological support (reference_map.py, topology.py), causal probabilistic occupancy and motion forecasting (forecasting.py), deterministic risk-, energy-, time-, progress-, and information-aware planning (planner.py), structured event telemetry (telemetry.py), and the integrated S4DTAMReference pipeline (pipeline.py).",
            "The reference implementation favors causality, deterministic replay, explicit uncertainty, and auditability over maximum throughput. This is advantageous during early numerical validation because an observed difference can be traced to a defined mechanism rather than hidden implementation nondeterminism.",
            "Several formulations in Sections 5.12-5.15 describe target research extensions rather than already completed backends. These include a learned hierarchical spatiotemporal model, a DBN/particle or equivalent multimodal probabilistic scene-state backend, an MPPI planner backend, adaptive octree/LOD storage, a closed-loop active-perception controller, and a safety-aware real-time scheduler. They are maintained as explicit implementation milestones and must preserve the same causal and evaluator-facing contracts when introduced.",
        ],
    )

    # Clarify implemented vs target mathematical backends without deleting the target equations.
    h512 = find_last(doc, "5.12. Bayesian Inference about Future Scene States")
    insert_after(
        h512,
        "Implementation status. The current executable reference uses a strictly causal probabilistic occupancy and motion forecaster driven only by observations available up to the prediction time. It produces Bernoulli occupancy probabilities, flow means, uncertainty estimates, observability masks, and physically timed forecast horizons, including irregular or dropped samples. This deterministic causal baseline is preferred for initial validation because future-frame leakage is structurally impossible. The DBN and particle approximation below define the target probabilistic extension for multimodal future-state uncertainty and are not claimed as the current default backend.",
    )

    h513 = find_last(doc, "5.13. Bayesian Risk Estimation and Trajectory Planning")
    insert_after(
        h513,
        "Implementation status. The executable reference planner currently uses deterministic finite-horizon beam search over admissible acceleration controls with explicit kinematic limits and decomposed collision-risk, energy, time, goal-progress, and information-value terms. This approach is more auditable than stochastic sampling for early confirmatory experiments because identical inputs generate the same candidate evaluation and decision trace. MPPI remains a planned backend for richer dynamics and will be introduced behind the same planner contract only after direct comparison with the deterministic reference.",
    )
    replace_contains(
        doc,
        "Since direct minimization of the expected value is computationally expensive, approximation by particle trajectory sampling via MPPI",
        "For a future high-dimensional stochastic planning backend, direct minimization of the expected value may be approximated by particle trajectory sampling via MPPI (Model Predictive Path Integral). The following formulation defines that target extension; it is not required by the current deterministic reference planner:",
    )

    h514 = find_last(doc, "5.14. Adaptive World Model with Hierarchical Resolution")
    insert_after(
        h514,
        "Implementation status. The current token memory already enforces bounded history and configurable limits on token count, map memory, and update time, together with lifecycle activation, sleeping, reactivation, merging, and removal rules. These mechanisms implement the resource-bounded principle of this section. The full adaptive octree and Level-of-Detail storage policy described below remains a planned world-model backend and will be validated separately before becoming the default representation.",
    )

    h515 = find_last(doc, "5.15. Computational Resource Management and Bottleneck Elimination")
    insert_after(
        h515,
        "Implementation status. Structured telemetry and resource budgets already expose update-time, token-count, memory, lifecycle, and decision events needed for profiling and replay. Predictive queue-overflow control, EDF-RT scheduling, guaranteed safety-critical CPU reservations, and the full survival-mode controller described below remain planned real-time systems extensions rather than features of the current Python reference.",
    )

    # Synchronize H1-H7 with the frozen preregistration.
    replace_section(
        doc,
        "6. Research Hypotheses",
        "7. Proposed Experimental Validation",
        [
            ("The confirmatory mechanism study uses seven preregistered hypotheses. Each hypothesis compares the full S4D-TAM model with a variant that disables exactly one mechanism while data, seeds, preprocessing, optimization budget, and all other settings remain fixed. External systems are deliberately excluded from this causal ablation family.", "Normal"),
            ("H1: Semantics improve navigation in scenarios containing dynamic objects. The primary endpoint is mission success; collisions per kilometer and semantic mIoU are secondary outcomes.", "Normal"),
            ("H2: Temporal state improves 1 s future-state prediction. The primary endpoint is forecast occupancy IoU at 1 s; flow EPE and temporal flip rate are secondary outcomes.", "Normal"),
            ("H3: Calibrated uncertainty reduces navigation risk. The primary endpoint is collisions per kilometer; ECE, NLL, and near misses are secondary outcomes.", "Normal"),
            ("H4: Topological support increases mission success. Path efficiency and ATE RMSE are secondary outcomes.", "Normal"),
            ("H5: Reference-map support reduces localization error. The primary endpoint is ATE RMSE; RPE and final drift percentage are secondary outcomes.", "Normal"),
            ("H6: Predictive risk modeling reduces collisions. The primary endpoint is collisions per kilometer; near misses and minimum clearance are secondary outcomes.", "Normal"),
            ("H7: Token lifecycle management reduces computational cost without materially degrading navigation quality. Confirmation requires both lower p95 latency and non-inferior mission success with a preregistered margin of 2 percentage points; map bytes and peak RSS are secondary outcomes.", "Normal"),
            ("H1-H6 are superiority hypotheses in the preregistered direction. H7 combines efficiency superiority with mission-success non-inferiority. The seven confirmatory decisions form one family and use the Holm procedure defined in Section 7.4.", "Normal"),
        ],
    )

    # Dataset protocol.
    replace_section(
        doc,
        "7.1. Datasets",
        "7.2. Two-Level Comparison Design",
        [
            ("The validation framework targets TartanAir, Blackbird, MARSIM, and AeroVerse, each converted to the common SequenceData contract with immutable source version, license record, coordinate convention, calibration metadata, and SHA-256 manifest. The cohort identifier frozen-2026-01 denotes the study cohort rather than a vendor release and must resolve to exact source versions and file hashes.", "Normal"),
            ("Data are split at the environment or flight level, never by adjacent frames. The train split is the only source of learned weight updates; calibration is reserved for thresholds, normalization, early stopping, and uncertainty calibration; test remains inaccessible until code, artifacts, and analysis are frozen.", "Normal"),
            ("The preregistered offline mechanism study uses 20 preselected test sequences per dataset, yielding 80 paired sequence-level units. Systems validation additionally defines 30 paired SIL scenarios with five seeds, 20 HIL scenarios with five seeds, and 20 paired controlled real-flight missions per variant when the safety gate remains open.", "Normal"),
            ("A dedicated UAV dataset remains desirable for synchronized RGB, thermal, IMU, GNSS, point cloud, elevation/reference map, static and dynamic labels, visibility or observability masks, and controlled sensor-degradation scenarios. Metrics that require unavailable ground truth are reported as unavailable rather than silently omitted or imputed.", "Normal"),
        ],
    )

    # Two-level comparison design with explicit subheadings.
    h72 = find_last(doc, "7.2. Two-Level Comparison Design")
    end73 = find_last(doc, "7.3. Metrics")
    remove_between(h72, end73)
    insert_blocks_before(
        end73,
        [
            ("The evaluation is split into two non-interchangeable levels. Both levels use common dataset and metric contracts where applicable, but they answer different scientific questions and produce separate statistical reports.", "Normal"),
            ("7.2.1. External System Comparison", "Heading 3"),
            ("The external comparison asks whether the complete S4D-TAM system is competitive with independently implemented navigation and SLAM systems under the same data, hardware policy, sensor availability, and evaluator definitions.", "Normal"),
            ("The mandatory primary baseline set comprises ORB-SLAM3, VINS-Mono, FAST-LIO2, and LIO-SAM. Each baseline is pinned to an upstream revision and fixed environment or container, calibration, sensor-topic mapping, loop-closure policy, warm-up policy, and hardware configuration. Results enter the benchmark only through the common AlgorithmResult artifact schema.", "Normal"),
            ("Additional external systems may be included only when their sensing assumptions are compatible, their exact implementation can be reproduced, and their output can be normalized without method-specific post-processing. Such additions must be frozen before the external study is executed.", "Normal"),
            ("This level measures complete-system competitiveness. A difference relative to an external baseline cannot be interpreted as the causal effect of an individual S4D-TAM component.", "Normal"),
            ("7.2.2. Internal Mechanism Study", "Heading 3"),
            ("The internal comparison tests mechanism contribution using the full system and seven single-component ablations: H1_no_semantics, H2_no_temporal_state, H3_no_calibrated_uncertainty, H4_no_topology, H5_no_reference_map, H6_no_risk_prediction, and H7_no_token_lifecycle.", "Normal"),
            ("Each variant differs from full by exactly one switch and uses its own frozen trained artifact. External systems never appear in this matrix. This separation prevents confounding complete-system identity with the causal effect of removing one S4D-TAM mechanism.", "Normal"),
        ],
    )

    # Metrics aligned with evaluator and preregistration.
    replace_section(
        doc,
        "7.3. Metrics",
        "8. Discussion",
        [
            ("Primary endpoints are deliberately compact and preregistered: localization uses ATE RMSE after SE(3) alignment without scale; forecasting uses occupancy IoU at 1 s and, for broader characterization, 3 s; safety uses mission success and collisions per kilometer; efficiency uses p95 latency and persistent map-memory footprint.", "Normal"),
            ("Localization and SLAM secondary metrics include RPE translation RMSE, final drift in meters and percentage of reference path length, and rotation RPE when orientation ground truth is available. Relocalization success rate, time to relocalize, and post-relocalization pose error are reported when the dataset supports those events.", "Normal"),
            ("Semantic metrics include mIoU and macro F1. Forecasting metrics include per-horizon occupancy IoU/F1, flow EPE, Brier score, NLL, and ECE. Pose uncertainty is evaluated by NEES, 95% coverage, and NLL where covariance output is available. Risk prediction may additionally report AUROC, false-alarm rate, and miss rate.", "Normal"),
            ("Navigation metrics include mission success, collisions per kilometer, near misses, minimum clearance, and path efficiency. Efficiency metrics include latency p95, map bytes, peak RSS, CPU time, and implementation-specific telemetry such as processed-token counts when consistently available.", "Normal"),
            ("Scale error, embedding quality, spatial-relation accuracy, emergency replanning counts, FPS, and energy consumption may be reported as secondary or exploratory metrics when their ground truth and measurement procedure are explicitly defined. They do not become confirmatory endpoints merely because an implementation emits them.", "Normal"),
            ("A missing metric caused by absent ground truth or absent algorithm output is recorded as unavailable. Algorithm failure, timeout, abort, or collision is not treated as a missing observation and remains visible in the safety analysis.", "Normal"),
            ("7.4. Confirmatory Protocol and Statistical Analysis", "Heading 2"),
            ("The preregistered H1-H7 study uses sequence or mission level inference rather than treating individual frames as independent samples. Continuous outcomes are summarized with paired mean differences and 95% BCa bootstrap confidence intervals using 10,000 resamples. Mission success is modeled as a paired scenario-level probability outcome, while distance-normalized event counts use Poisson or negative-binomial modeling with log distance as an offset according to the preregistered overdispersion rule.", "Normal"),
            ("The seven H1-H7 confirmatory decisions form one family and are corrected with the sequential Holm procedure at family-wise alpha = 0.05. H1-H6 require the corrected result and confidence interval to support the predicted superiority direction. H7 additionally requires mission-success non-inferiority within 2 percentage points while improving p95 latency.", "Normal"),
            ("No missing technical result is imputed. Aborts and algorithm failures remain in the denominator for safety outcomes; complete-case reporting is accompanied by the preregistered worst-case sensitivity analysis. Exclusions are limited to integrity failures, excessive missing timestamps, independently confirmed sensor faults, or operational constraints defined before result inspection.", "Normal"),
            ("The external system comparison is reported separately from this H1-H7 inferential family. It uses the same frozen datasets and common evaluators, but its effect estimates describe differences between complete systems rather than component-level causal effects.", "Normal"),
        ],
    )

    # Discussion and conclusion no longer describe prototype implementation as future work.
    replace_contains(
        doc,
        "Future work should focus on implementing a system prototype",
        "Future work should focus on completing and training the learned S4D-TAM model, validating public-dataset converters and pinned external baselines, executing the preregistered H1-H7 mechanism study, and progressing from offline evaluation through SIL, HIL, and controlled real-flight validation. Additional engineering work includes the probabilistic DBN/particle backend, MPPI planning, adaptive octree/LOD storage, active perception, and safety-aware real-time scheduling.",
    )

    replace_contains(
        doc,
        "1. Prototype implementation of the S4D-TAM system and experimental validation on the TartanAir",
        "1. Completion and training of the learned S4D-TAM model followed by two-level validation: external system comparison on TartanAir, Blackbird, MARSIM, and AeroVerse, and the preregistered internal full-vs-H1-H7 mechanism study.",
    )

    replace_contains(
        doc,
        "Seven research hypotheses have been formulated regarding computational efficiency",
        "Seven preregistered mechanism hypotheses have been formulated for semantics, temporal state, calibrated uncertainty, topology, reference-map support, predictive risk, and token lifecycle. Their causal full-vs-H1-H7 study is explicitly separated from the external system comparison against ORB-SLAM3, VINS-Mono, FAST-LIO2, and LIO-SAM.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
