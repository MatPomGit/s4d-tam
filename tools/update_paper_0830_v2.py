from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


MODULE_PATH = Path(__file__).with_name("update_paper_0830.py")
OUTPUT_PATH = Path("paper/S4D_TAM_draft_0830_updated.docx")
spec = importlib.util.spec_from_file_location("update_paper_0830", MODULE_PATH)
if spec is None or spec.loader is None:
    raise RuntimeError(f"Cannot load paper updater from {MODULE_PATH}")
updater = importlib.util.module_from_spec(spec)
spec.loader.exec_module(updater)


def _normalize(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").replace("\u2011", "-").split())


def robust_replace_exact(doc, old: str, new: str, occurrence: str = "last") -> None:
    matches = [p for p in doc.paragraphs if p.text.strip() == old]
    if not matches:
        raise ValueError(f"Paragraph not found: {old}")
    if old == "7.2. System Variants":
        for paragraph in matches:
            paragraph.text = new
        return
    paragraph = matches[0] if occurrence == "first" else matches[-1]
    paragraph.text = new


def robust_replace_contains(doc, needle: str, new: str) -> None:
    wanted = _normalize(needle)
    for element in doc.element.iter(qn("w:p")):
        paragraph = Paragraph(element, doc)
        if wanted in _normalize(paragraph.text):
            paragraph.text = new
            return
    raise ValueError(f"Text fragment not found after normalization: {needle}")


def _paragraphs(doc: Document, text: str) -> list[Paragraph]:
    return [paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == text]


def _format_body_heading(paragraph: Paragraph, *, size_pt: float | None) -> None:
    paragraph.style = "Normal"
    paragraph.paragraph_format.left_indent = None
    for run in paragraph.runs:
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt) if size_pt is not None else None


def finalize_layout() -> None:
    doc = Document(OUTPUT_PATH)

    toc_indents = {
        "4.3. Current Reference Implementation and Target Research Modules": 0.25,
        "7.2.1. External System Comparison": 0.50,
        "7.2.2. Internal Mechanism Study": 0.50,
        "7.4. Confirmatory Protocol and Statistical Analysis": 0.25,
    }
    for text, indent_in in toc_indents.items():
        matches = _paragraphs(doc, text)
        if len(matches) < 2:
            raise ValueError(f"Expected TOC and body occurrences for: {text}")
        matches[0].style = "Normal"
        matches[0].paragraph_format.left_indent = Inches(indent_in)
        for run in matches[0].runs:
            run.bold = None
            run.font.name = "Times New Roman"
            run.font.size = None

    body_headings = {
        "4.3. Current Reference Implementation and Target Research Modules": 12.0,
        "7.2.1. External System Comparison": None,
        "7.2.2. Internal Mechanism Study": None,
        "7.4. Confirmatory Protocol and Statistical Analysis": 12.0,
    }
    for text, size_pt in body_headings.items():
        _format_body_heading(_paragraphs(doc, text)[-1], size_pt=size_pt)

    doc.save(OUTPUT_PATH)


updater.replace_exact = robust_replace_exact
updater.replace_contains = robust_replace_contains
updater.main()
finalize_layout()
